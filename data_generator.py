"""
Synthetic Contract Dataset Generator for CLM Automation System.
Generates 10-15 diverse contract documents with variations, conflicts, and metadata.
"""
import os
import random
from datetime import datetime, timedelta
import json

# Try to import docx, fall back to simple text if not available
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Word documents will be created as text files.")

class ContractDataGenerator:
    def __init__(self, output_dir="./documents"):
        self.output_dir = output_dir
        self.companies = [
            "TechCorp Solutions", "Global Industries Ltd", "Innovation Partners",
            "Digital Dynamics Inc", "Future Systems LLC", "Smart Technologies Corp"
        ]
        self.contract_types = [
            "Service Agreement", "Software License", "Consulting Contract",
            "Non-Disclosure Agreement", "Employment Contract", "Vendor Agreement"
        ]
        
        # Create conflicting data for testing
        self.conflicting_addresses = {
            "TechCorp Solutions": ["123 Main St, New York, NY 10001", "456 Business Ave, New York, NY 10002"],
            "Global Industries Ltd": ["789 Corporate Blvd, Los Angeles, CA 90210", "321 Enterprise Way, Los Angeles, CA 90211"]
        }
        
        self.conflicting_dates = {
            "contract_001": ["2024-12-31", "2025-01-15"],
            "contract_002": ["2024-11-30", "2024-12-15"]
        }

    def generate_pdf_contracts(self):
        """Generate 4-5 PDF contracts (simulated as text files for this demo)"""
        contracts = [
            {
                "filename": "service_agreement_v1.txt",
                "content": self._create_service_agreement_v1(),
                "metadata": {"type": "service_agreement", "version": "1.0", "department": "legal"}
            },
            {
                "filename": "software_license_v2.txt", 
                "content": self._create_software_license_v2(),
                "metadata": {"type": "software_license", "version": "2.0", "department": "it"}
            },
            {
                "filename": "consulting_contract_v1.txt",
                "content": self._create_consulting_contract_v1(),
                "metadata": {"type": "consulting", "version": "1.0", "department": "operations"}
            },
            {
                "filename": "nda_agreement_v1.txt",
                "content": self._create_nda_agreement_v1(),
                "metadata": {"type": "nda", "version": "1.0", "department": "legal"}
            },
            {
                "filename": "vendor_agreement_v1.txt",
                "content": self._create_vendor_agreement_v1(),
                "metadata": {"type": "vendor", "version": "1.0", "department": "procurement"}
            }
        ]
        
        for contract in contracts:
            filepath = os.path.join(self.output_dir, "pdfs", contract["filename"])
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(contract["content"])
            
            # Save metadata
            metadata_path = filepath.replace('.txt', '_metadata.json')
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(contract["metadata"], f, indent=2)

    def generate_word_contracts(self):
        """Generate 3-4 Word documents"""
        contracts = [
            {
                "filename": "employment_contract_draft.docx",
                "content": self._create_employment_contract(),
                "metadata": {"type": "employment", "status": "draft", "department": "hr"}
            },
            {
                "filename": "contract_amendment_001.docx",
                "content": self._create_contract_amendment(),
                "metadata": {"type": "amendment", "original_contract": "service_agreement_v1", "department": "legal"}
            },
            {
                "filename": "renewal_agreement_2024.docx",
                "content": self._create_renewal_agreement(),
                "metadata": {"type": "renewal", "year": "2024", "department": "legal"}
            },
            {
                "filename": "termination_notice.docx",
                "content": self._create_termination_notice(),
                "metadata": {"type": "termination", "status": "active", "department": "legal"}
            }
        ]
        
        for contract in contracts:
            if DOCX_AVAILABLE:
                doc = Document()
                doc.add_heading(contract["content"]["title"], 0)
                
                for section in contract["content"]["sections"]:
                    doc.add_heading(section["heading"], level=1)
                    doc.add_paragraph(section["content"])
                
                filepath = os.path.join(self.output_dir, "word", contract["filename"])
                doc.save(filepath)
            else:
                # Create as text file instead
                content = contract["content"]["title"] + "\n\n"
                for section in contract["content"]["sections"]:
                    content += f"{section['heading']}\n"
                    content += f"{section['content']}\n\n"
                
                filepath = os.path.join(self.output_dir, "word", contract["filename"].replace('.docx', '.txt'))
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            # Save metadata
            metadata_path = filepath.replace('.docx', '_metadata.json').replace('.txt', '_metadata.json')
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(contract["metadata"], f, indent=2)

    def generate_text_files(self):
        """Generate 2-3 text files (summaries, emails)"""
        text_files = [
            {
                "filename": "contract_summary_2024.txt",
                "content": self._create_contract_summary(),
                "metadata": {"type": "summary", "year": "2024", "department": "legal"}
            },
            {
                "filename": "email_correspondence_contract.txt",
                "content": self._create_email_correspondence(),
                "metadata": {"type": "email", "subject": "contract_discussion", "department": "legal"}
            },
            {
                "filename": "contract_negotiation_notes.txt",
                "content": self._create_negotiation_notes(),
                "metadata": {"type": "notes", "meeting": "contract_negotiation", "department": "legal"}
            }
        ]
        
        for file in text_files:
            filepath = os.path.join(self.output_dir, "text", file["filename"])
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(file["content"])
            
            # Save metadata
            metadata_path = filepath.replace('.txt', '_metadata.json')
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(file["metadata"], f, indent=2)

    def generate_unstructured_files(self):
        """Generate 2 unstructured files (meeting notes, etc.)"""
        unstructured_files = [
            {
                "filename": "meeting_notes_contract_review.txt",
                "content": self._create_meeting_notes(),
                "metadata": {"type": "meeting_notes", "date": "2024-01-15", "department": "legal"}
            },
            {
                "filename": "contract_issues_discussion.txt",
                "content": self._create_issues_discussion(),
                "metadata": {"type": "discussion", "topic": "contract_issues", "department": "legal"}
            }
        ]
        
        for file in unstructured_files:
            filepath = os.path.join(self.output_dir, "unstructured", file["filename"])
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(file["content"])
            
            # Save metadata
            metadata_path = filepath.replace('.txt', '_metadata.json')
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(file["metadata"], f, indent=2)

    def _create_service_agreement_v1(self):
        """Create Service Agreement v1 with conflicting information"""
        return f"""
SERVICE AGREEMENT

Contract ID: SA-2024-001
Version: 1.0
Effective Date: January 1, 2024
Expiration Date: December 31, 2024

PARTIES:
Client: TechCorp Solutions
Address: {self.conflicting_addresses['TechCorp Solutions'][0]}
Contact: John Smith, CEO
Email: john.smith@techcorp.com

Service Provider: Global Industries Ltd
Address: {self.conflicting_addresses['Global Industries Ltd'][0]}
Contact: Sarah Johnson, VP Operations
Email: s.johnson@globalindustries.com

SCOPE OF SERVICES:
The Service Provider agrees to provide comprehensive IT consulting services including:
- System architecture design
- Implementation support
- 24/7 technical support
- Monthly reporting

TERMS AND CONDITIONS:
1. Payment Terms: Net 30 days
2. Service Level Agreement: 99.9% uptime guarantee
3. Confidentiality: Both parties agree to maintain strict confidentiality
4. Termination: Either party may terminate with 30 days written notice
5. Governing Law: State of New York

COMPENSATION:
Monthly retainer: $15,000
Additional services: $150/hour
Maximum monthly cap: $25,000

SIGNATURES:
Client: _________________ Date: _________
Service Provider: _________________ Date: _________

This agreement shall be binding upon execution by both parties.
"""

    def _create_software_license_v2(self):
        """Create Software License v2 with different expiration date"""
        return f"""
SOFTWARE LICENSE AGREEMENT

Contract ID: SLA-2024-002
Version: 2.0
Effective Date: February 1, 2024
Expiration Date: {self.conflicting_dates['contract_001'][0]}

LICENSOR: Innovation Partners
Address: 555 Tech Street, San Francisco, CA 94105
Contact: Michael Chen, CTO

LICENSEE: Digital Dynamics Inc
Address: 888 Innovation Drive, Austin, TX 73301
Contact: Lisa Rodriguez, IT Director

LICENSED SOFTWARE:
"SmartAnalytics Pro" - Business Intelligence Platform
Version: 3.2.1
License Type: Enterprise Multi-User

LICENSE TERMS:
1. License Grant: Non-exclusive, non-transferable license
2. User Limit: Up to 100 concurrent users
3. Deployment: On-premises and cloud environments
4. Updates: Included for 12 months from effective date
5. Support: 24/7 technical support included

RESTRICTIONS:
- No reverse engineering
- No redistribution without written consent
- No use beyond licensed scope

FEES:
Annual License Fee: $50,000
Support and Maintenance: $10,000/year
Implementation Services: $25,000 (one-time)

AUTO-RENEWAL:
This license shall automatically renew for successive one-year terms unless terminated with 90 days written notice.

SIGNATURES:
Licensor: _________________ Date: _________
Licensee: _________________ Date: _________
"""

    def _create_consulting_contract_v1(self):
        """Create Consulting Contract with approaching expiration"""
        exp_date = (datetime.now() + timedelta(days=25)).strftime("%B %d, %Y")
        return f"""
CONSULTING SERVICES AGREEMENT

Contract ID: CSA-2024-003
Effective Date: March 1, 2024
Expiration Date: {exp_date}

CONSULTANT: Future Systems LLC
Address: 777 Business Park, Chicago, IL 60601
Contact: David Wilson, Principal Consultant
Phone: (312) 555-0123

CLIENT: Smart Technologies Corp
Address: 999 Corporate Plaza, Boston, MA 02101
Contact: Jennifer Lee, VP Engineering
Phone: (617) 555-0456

SCOPE OF WORK:
The Consultant shall provide the following services:
1. Technology assessment and recommendations
2. Process optimization consulting
3. Team training and development
4. Implementation planning and oversight

DELIVERABLES:
- Comprehensive assessment report (due within 30 days)
- Process improvement recommendations
- Training materials and sessions
- Implementation roadmap

COMPENSATION:
Project Fee: $75,000 (fixed price)
Expenses: Reimbursed up to $5,000
Payment Schedule: 50% upon contract execution, 50% upon completion

TERMINATION:
Either party may terminate this agreement with 15 days written notice.
Upon termination, Consultant shall be paid for work completed to date.

CONFIDENTIALITY:
All information exchanged shall be treated as confidential and proprietary.

SIGNATURES:
Consultant: _________________ Date: _________
Client: _________________ Date: _________
"""

    def _create_nda_agreement_v1(self):
        """Create NDA Agreement"""
        return f"""
NON-DISCLOSURE AGREEMENT

Contract ID: NDA-2024-004
Effective Date: April 1, 2024
Expiration Date: April 1, 2026

DISCLOSING PARTY: TechCorp Solutions
Address: {self.conflicting_addresses['TechCorp Solutions'][1]}
Contact: John Smith, CEO

RECEIVING PARTY: Innovation Partners
Address: 555 Tech Street, San Francisco, CA 94105
Contact: Michael Chen, CTO

PURPOSE:
The parties wish to explore potential business opportunities and may share confidential information.

CONFIDENTIAL INFORMATION:
Includes but not limited to:
- Technical specifications and designs
- Business plans and strategies
- Financial information
- Customer lists and data
- Proprietary processes and methods

OBLIGATIONS:
The Receiving Party agrees to:
1. Hold all Confidential Information in strict confidence
2. Use Confidential Information solely for the stated purpose
3. Not disclose to third parties without written consent
4. Return or destroy all materials upon request

TERM:
This agreement shall remain in effect for 2 years from the effective date.

GOVERNING LAW:
This agreement shall be governed by the laws of California.

SIGNATURES:
Disclosing Party: _________________ Date: _________
Receiving Party: _________________ Date: _________
"""

    def _create_vendor_agreement_v1(self):
        """Create Vendor Agreement"""
        return f"""
VENDOR SERVICES AGREEMENT

Contract ID: VSA-2024-005
Effective Date: May 1, 2024
Expiration Date: May 1, 2025

VENDOR: Global Industries Ltd
Address: {self.conflicting_addresses['Global Industries Ltd'][1]}
Contact: Sarah Johnson, VP Operations
Tax ID: 12-3456789

CUSTOMER: Digital Dynamics Inc
Address: 888 Innovation Drive, Austin, TX 73301
Contact: Lisa Rodriguez, Procurement Manager

SERVICES:
The Vendor shall provide:
1. Office supplies and equipment
2. IT hardware and software
3. Maintenance and support services
4. Delivery and installation

PRICING:
- Office Supplies: 15% discount from retail
- IT Equipment: 20% discount from MSRP
- Services: $125/hour
- Delivery: Free for orders over $500

PAYMENT TERMS:
Net 30 days from invoice date
Late payment fee: 1.5% per month

QUALITY STANDARDS:
All products must meet industry standards
Warranty: 1 year on all products
Support: 8x5 technical support

TERMINATION:
Either party may terminate with 60 days written notice.

SIGNATURES:
Vendor: _________________ Date: _________
Customer: _________________ Date: _________
"""

    def _create_employment_contract(self):
        """Create Employment Contract Word document"""
        return {
            "title": "EMPLOYMENT AGREEMENT",
            "sections": [
                {
                    "heading": "Employee Information",
                    "content": "Name: Alex Thompson\nPosition: Senior Software Engineer\nDepartment: Engineering\nStart Date: June 1, 2024\nEmployee ID: EMP-2024-001"
                },
                {
                    "heading": "Company Information",
                    "content": "Company: Smart Technologies Corp\nAddress: 999 Corporate Plaza, Boston, MA 02101\nContact: Jennifer Lee, VP Engineering"
                },
                {
                    "heading": "Compensation",
                    "content": "Base Salary: $120,000 annually\nBonus: Up to 20% based on performance\nBenefits: Health, dental, vision, 401k matching\nVacation: 20 days annually"
                },
                {
                    "heading": "Terms and Conditions",
                    "content": "Employment Type: Full-time, at-will\nProbation Period: 90 days\nConfidentiality: Required per company policy\nNon-compete: 12 months after termination\nNotice Period: 2 weeks for resignation"
                }
            ]
        }

    def _create_contract_amendment(self):
        """Create Contract Amendment Word document"""
        return {
            "title": "CONTRACT AMENDMENT #1",
            "sections": [
                {
                    "heading": "Original Contract",
                    "content": "Service Agreement SA-2024-001\nEffective Date: January 1, 2024\nParties: TechCorp Solutions and Global Industries Ltd"
                },
                {
                    "heading": "Amendment Details",
                    "content": "Amendment Date: July 1, 2024\nReason: Scope expansion and pricing adjustment\nApproved by: Legal Department"
                },
                {
                    "heading": "Changes",
                    "content": "1. Monthly retainer increased from $15,000 to $18,000\n2. Additional services rate increased from $150/hour to $175/hour\n3. New service: Cloud migration support\n4. Extended term by 6 months to June 30, 2025"
                },
                {
                    "heading": "Signatures",
                    "content": "Client: _________________ Date: _________\nService Provider: _________________ Date: _________"
                }
            ]
        }

    def _create_renewal_agreement(self):
        """Create Renewal Agreement Word document"""
        return {
            "title": "CONTRACT RENEWAL AGREEMENT 2024",
            "sections": [
                {
                    "heading": "Renewal Information",
                    "content": "Original Contract: Software License SLA-2024-002\nRenewal Period: January 1, 2025 - December 31, 2025\nRenewal Date: November 15, 2024"
                },
                {
                    "heading": "Updated Terms",
                    "content": "License Fee: $55,000 (10% increase)\nUser Limit: Increased to 150 concurrent users\nSupport: Extended to 18x7 coverage\nNew Features: Advanced analytics module included"
                },
                {
                    "heading": "Conditions",
                    "content": "Renewal subject to:\n1. Satisfactory performance during previous term\n2. Payment of all outstanding invoices\n3. Execution of updated terms and conditions"
                }
            ]
        }

    def _create_termination_notice(self):
        """Create Termination Notice Word document"""
        return {
            "title": "CONTRACT TERMINATION NOTICE",
            "sections": [
                {
                    "heading": "Termination Details",
                    "content": "Contract: Vendor Services Agreement VSA-2024-005\nTermination Date: October 31, 2024\nNotice Date: September 1, 2024\nReason: Vendor performance issues"
                },
                {
                    "heading": "Transition Plan",
                    "content": "1. All outstanding orders to be completed by termination date\n2. Return of all company property within 30 days\n3. Final payment within 45 days of termination\n4. New vendor selection process initiated"
                },
                {
                    "heading": "Contact Information",
                    "content": "For questions regarding this termination:\nContact: Lisa Rodriguez, Procurement Manager\nEmail: l.rodriguez@digitaldynamics.com\nPhone: (512) 555-0789"
                }
            ]
        }

    def _create_contract_summary(self):
        """Create contract summary text file"""
        return """
CONTRACT PORTFOLIO SUMMARY - 2024

ACTIVE CONTRACTS (12):
1. Service Agreement SA-2024-001 - TechCorp Solutions - $15,000/month
2. Software License SLA-2024-002 - Innovation Partners - $50,000/year
3. Consulting Contract CSA-2024-003 - Future Systems LLC - $75,000 project
4. NDA Agreement NDA-2024-004 - Innovation Partners - No cost
5. Vendor Agreement VSA-2024-005 - Global Industries Ltd - Variable
6. Employment Contract EMP-2024-001 - Alex Thompson - $120,000/year

EXPIRING SOON (Next 30 days):
- Consulting Contract CSA-2024-003 (expires in 25 days)
- Software License SLA-2024-002 (expires in 15 days)

RENEWALS PENDING:
- Service Agreement SA-2024-001 (renewal decision due by Nov 30)
- Vendor Agreement VSA-2024-005 (renewal decision due by Dec 15)

CONFLICTS IDENTIFIED:
1. TechCorp Solutions has different addresses in contracts SA-2024-001 and NDA-2024-004
2. Global Industries Ltd address inconsistency between VSA-2024-005 and other contracts
3. Contract SLA-2024-002 has conflicting expiration dates in different versions

TOTAL CONTRACT VALUE: $1,250,000 annually
RISK LEVEL: Medium (2 contracts expiring soon, 3 conflicts identified)
"""

    def _create_email_correspondence(self):
        """Create email correspondence text file"""
        return """
From: john.smith@techcorp.com
To: s.johnson@globalindustries.com
Subject: Contract Amendment Discussion
Date: June 15, 2024

Hi Sarah,

I hope this email finds you well. I wanted to follow up on our discussion about the service agreement amendment.

As we discussed in our meeting last week, we need to address a few items:

1. The monthly retainer increase from $15,000 to $18,000
2. The new cloud migration services you'll be providing
3. The extended contract term through June 2025

I've attached the draft amendment for your review. Please let me know if you have any questions or concerns.

Also, I noticed there's a discrepancy in our address information. Our legal team wants to ensure we have the correct address on file. Could you confirm which address should be used for official correspondence?

Looking forward to your response.

Best regards,
John Smith
CEO, TechCorp Solutions

---
From: s.johnson@globalindustries.com
To: john.smith@techcorp.com
Subject: Re: Contract Amendment Discussion
Date: June 16, 2024

Hi John,

Thanks for the follow-up. I've reviewed the amendment and everything looks good to me.

Regarding the address discrepancy, I apologize for the confusion. Our official business address is 456 Business Ave, New York, NY 10002. The other address in some of our contracts is our old location that we moved from last year.

I'll have our legal team update all our contracts with the correct address information.

The amendment looks good to proceed. When would you like to schedule the signing?

Best,
Sarah Johnson
VP Operations, Global Industries Ltd
"""

    def _create_negotiation_notes(self):
        """Create negotiation notes text file"""
        return """
CONTRACT NEGOTIATION NOTES
Meeting Date: March 20, 2024
Participants: John Smith (TechCorp), Sarah Johnson (Global Industries), Legal Teams

AGENDA ITEMS DISCUSSED:

1. PRICING ADJUSTMENTS
- Current monthly retainer: $15,000
- Proposed increase: $18,000 (20% increase)
- Rationale: Increased service scope and inflation adjustment
- Counter-proposal: $16,500 (10% increase)
- Resolution: Agreed on $17,000 (13% increase)

2. SERVICE SCOPE EXPANSION
- New services requested: Cloud migration support
- Additional hours: 20 hours/month
- Rate for new services: $175/hour
- Implementation timeline: 30 days after contract execution

3. CONTRACT TERM
- Current expiration: December 31, 2024
- Proposed extension: June 30, 2025
- Rationale: Align with fiscal year and project timeline
- Agreement: Extended to June 30, 2025

4. ADDRESS DISCREPANCY
- Issue: Different addresses in various contracts
- TechCorp address: 123 Main St vs 456 Business Ave
- Global Industries address: 789 Corporate Blvd vs 321 Enterprise Way
- Action: Both parties to provide official addresses for contract updates

5. NEXT STEPS
- Legal teams to draft amendment
- Address verification and updates
- Contract signing scheduled for April 1, 2024
- Implementation to begin April 15, 2024

NOTES:
- Both parties committed to long-term partnership
- Discussion of potential additional services in Q3 2024
- Confidentiality maintained throughout negotiation
"""

    def _create_meeting_notes(self):
        """Create unstructured meeting notes"""
        return """
Contract Review Meeting - Jan 15, 2024

Attendees: Legal team, Procurement, IT dept

Quick notes from the meeting:

- Need to review all contracts expiring this year
- Some contracts have conflicting info (addresses, dates)
- TechCorp contract needs renewal decision by Nov
- Global Industries contract has performance issues
- New vendor evaluation process started

Action items:
- Sarah to check address discrepancies
- Mike to prepare renewal recommendations
- Lisa to start vendor search process

Issues found:
- Contract SA-2024-001 has wrong address for TechCorp
- Some contracts missing metadata
- Expiration dates not consistent across versions

Next meeting: Feb 15, 2024

Random thoughts:
- Should we implement automated contract monitoring?
- Need better version control for contract documents
- Consider using AI for contract analysis
"""

    def _create_issues_discussion(self):
        """Create unstructured issues discussion"""
        return """
Contract Issues Discussion - Internal Notes

Problems we're facing:

1. Address conflicts everywhere
- TechCorp has 2 different addresses in our system
- Global Industries also has conflicting addresses
- This is causing confusion in billing and legal docs

2. Date inconsistencies
- Some contracts show different expiration dates
- Not sure which version is correct
- Need to verify with original signed copies

3. Missing metadata
- Half our contracts don't have proper metadata
- Makes searching and filtering difficult
- Need to standardize this going forward

4. Renewal tracking
- Manual process is error-prone
- Missed a renewal deadline last quarter
- Need automated alerts

5. Version control nightmare
- Multiple versions of same contract
- Not clear which is current
- Need better document management

Solutions to consider:
- Implement contract management system
- Use AI to extract and standardize data
- Set up automated renewal alerts
- Create single source of truth for contract data

Priority: HIGH - These issues are affecting business operations
"""

    def generate_all_documents(self):
        """Generate all synthetic documents"""
        print("Generating synthetic contract dataset...")
        
        # Create subdirectories
        os.makedirs(os.path.join(self.output_dir, "pdfs"), exist_ok=True)
        os.makedirs(os.path.join(self.output_dir, "word"), exist_ok=True)
        os.makedirs(os.path.join(self.output_dir, "text"), exist_ok=True)
        os.makedirs(os.path.join(self.output_dir, "unstructured"), exist_ok=True)
        
        # Generate documents
        self.generate_pdf_contracts()
        self.generate_word_contracts()
        self.generate_text_files()
        self.generate_unstructured_files()
        
        print(f"Generated synthetic dataset in {self.output_dir}")
        print("Documents created:")
        print("- 5 PDF contracts (simulated as text)")
        print("- 4 Word documents")
        print("- 3 Text files")
        print("- 2 Unstructured files")
        print("Total: 14 documents with variations, conflicts, and metadata")

if __name__ == "__main__":
    generator = ContractDataGenerator()
    generator.generate_all_documents()

